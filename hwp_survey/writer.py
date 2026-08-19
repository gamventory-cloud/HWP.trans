# -*- coding: utf-8 -*-
"""렌더링 블록 -> 워드(.docx) 설문지.

한글 문서에서 python-docx가 물어뜯기 쉬운 지점들:
  * run.font.name만 지정하면 한글에는 적용되지 않는다 -> w:rFonts의 eastAsia 필요
  * 표는 tblGrid(열 폭)와 각 셀 width를 둘 다 지정 + tblLayout=fixed
  * 셀 음영은 w:shd val="clear" (solid는 검게 렌더링)
  * 공백만 있는 문단은 렌더러가 잘라내므로 밑줄 칸은 nbsp 사용
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DEFAULTS = {
    "font": "맑은 고딕",
    "latin_font": "Arial",
    "base_pt": 10.5,
    "accent": "1F3B63",
    "single_mark": "○",
    "multi_mark": "□",
    "cell_mark": "□",
    "content_cm": 16.6,
    "row_label_cm": 9.0,
}


class SurveyWriter:
    def __init__(self, **opts):
        self.o = {**DEFAULTS, **{k: v for k, v in opts.items() if v is not None}}
        self.doc = self._new_document()
        self.qno = 0

    # ------------------------------------------------------------ 기본
    def _new_document(self):
        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = self.o["latin_font"]
        normal.font.size = Pt(self.o["base_pt"])
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), self.o["font"])
        normal.paragraph_format.space_after = Pt(4)
        normal.paragraph_format.line_spacing = 1.3

        sec = doc.sections[0]
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.2)
        return doc

    def font(self, run, size=None, bold=None, color=None, underline=None):
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), self.o["font"])       # 한글
        rFonts.set(qn("w:ascii"), self.o["latin_font"])    # 영문/숫자
        rFonts.set(qn("w:hAnsi"), self.o["latin_font"])
        if size:
            run.font.size = size
        if bold is not None:
            run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if underline is not None:
            run.font.underline = underline
        return run

    def para(self, text="", size=None, bold=False, align=None, indent=0.0,
             before=0, after=4, color=None):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(indent)
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        if align is not None:
            p.alignment = align
        if text:
            self.font(p.add_run(text), size=size or Pt(self.o["base_pt"]),
                      bold=bold, color=color)
        return p

    @staticmethod
    def border(paragraph, edges=("bottom",), sz=6, color="808080"):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        for edge in edges:
            el = pBdr.makeelement(qn(f"w:{edge}"), {})
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "3")
            el.set(qn("w:color"), color)
            pBdr.append(el)
        style = pPr.find(qn("w:pStyle"))
        pPr.insert(1 if style is not None else 0, pBdr)   # 스키마 순서 준수

    @staticmethod
    def shade(cell, fill="F2F2F2"):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"), {})
        shd.set(qn("w:val"), "clear")              # solid 금지
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    @staticmethod
    def fix_layout(table, header_repeat=True):
        tblPr = table._tbl.tblPr
        layout = tblPr.makeelement(qn("w:tblLayout"), {})
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        for i, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))
            if header_repeat and i == 0:
                trPr.append(trPr.makeelement(qn("w:tblHeader"), {}))

    def answer_line(self, indent=0.8, length=70, after=6):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(after)
        self.font(p.add_run("\u00a0" * length), underline=True)
        return p

    # ------------------------------------------------------------ 블록
    def write(self, blocks):
        for b in blocks:
            getattr(self, f"_{b['kind']}", self._note)(b)
        self.para("― 설문에 응해 주셔서 감사합니다. ―",
                  align=WD_ALIGN_PARAGRAPH.CENTER, before=16)
        return self

    def _title(self, b):
        p = self.para(b["text"], size=Pt(16), bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
        self.border(p)

    def _intro(self, b):
        self.para(b["text"], size=Pt(10), after=6)

    def _box(self, b):
        p = self.para(b["text"], size=Pt(9.5), indent=0.2, before=4, after=8)
        self.border(p, edges=("top", "bottom", "left", "right"), sz=4,
                    color="AAAAAA")

    def _section(self, b):
        p = self.para(b["text"], size=Pt(12), bold=True, before=14, after=8,
                      color=self.o["accent"])
        self.border(p, sz=12, color=self.o["accent"])

    def _note(self, b):
        self.para(b["text"], size=Pt(9), indent=0.6, color="606060")

    def _question(self, b):
        self.qno += 1
        suffix = " (복수 응답 가능)" if b["type"] == "복수" else ""
        self.para(f"{self.qno}. {b['text']}{suffix}", bold=True, before=10, after=5)
        for note in b.get("notes", []):
            self.para(note, size=Pt(9), indent=0.6, after=4, color="606060")

        if b["type"] == "표":
            self._matrix(b)
        elif b["type"] == "척도":
            self._scale(b)
        elif b["type"] == "장문":
            for _ in range(3):
                self.answer_line(after=10)
        elif b["type"] == "단답" and not b["options"]:
            self.answer_line()
        else:
            mark = self.o["multi_mark"] if b["type"] == "복수" else self.o["single_mark"]
            for opt in b["options"]:
                if opt["type"] == "group":
                    self.para(opt["text"], size=Pt(9.5), bold=True, indent=0.4,
                              before=4, after=2)
                else:
                    self.para(f"{mark} {opt['text']}", indent=0.8, after=2)
            if not b["options"]:
                self.answer_line()

    def _matrix(self, b):
        head = b["matrix"] or ["①", "②", "③", "④", "⑤"]
        rows = b["options"] or [{"type": "row", "text": "항목"}]
        first = Cm(min(self.o["row_label_cm"], self.o["content_cm"] - 1))
        rest = Cm((self.o["content_cm"] - first.cm) / len(head))

        table = self.doc.add_table(rows=1, cols=len(head) + 1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for i, label in enumerate(["구분"] + head):
            cell = table.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            self.font(p.add_run(label), size=Pt(8.5), bold=True)
            self.shade(cell)

        for item in rows:
            cells = table.add_row().cells
            if item["type"] == "group":                    # 표 안 소제목 행
                merged = cells[0].merge(cells[-1])
                p = merged.paragraphs[0]
                self.font(p.add_run(item["text"]), size=Pt(9), bold=True)
                self.shade(merged, "EAEFF7")
                continue
            p = cells[0].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            self.font(p.add_run(item["text"]), size=Pt(9.5))
            for i in range(1, len(head) + 1):
                cp = cells[i].paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.font(cp.add_run(self.o["cell_mark"]), size=Pt(11))

        widths = [first] + [rest] * len(head)
        for i, col in enumerate(table.columns):
            col.width = widths[i]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(widths):
                    cell.width = widths[i]
        self.fix_layout(table)
        self.para("", after=6)

    def _scale(self, b):
        lo, hi = b["scale"] or (1, 5)
        labels = [str(n) for n in range(lo, hi + 1)]
        table = self.doc.add_table(rows=2, cols=len(labels))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        w = Cm(min(12.0, self.o["content_cm"]) / len(labels))
        for col in table.columns:
            col.width = w
        for r, values in enumerate([labels, [self.o["single_mark"]] * len(labels)]):
            for c, value in enumerate(values):
                cell = table.rows[r].cells[c]
                cell.width = w
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.font(p.add_run(value), size=Pt(10), bold=(r == 0))
                if r == 0:
                    self.shade(cell)
        self.fix_layout(table, header_repeat=False)
        self.para("", after=4)

    # ------------------------------------------------------------ 출력
    def save(self, path):
        self.doc.save(path)
        return path

    def to_bytes(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()


def build_docx(blocks, **opts) -> bytes:
    return SurveyWriter(**opts).write(blocks).to_bytes()
